"""
Word Document Generator for College Memory Book
Generates a .docx file in Arabic (RTL) with meetings, colleagues, and memories.
Optimized for performance with large datasets.
"""
import os
import gc
import logging
from functools import lru_cache
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io
from django.conf import settings

logger = logging.getLogger(__name__)

# Global settings for image optimization
MAX_IMAGE_WIDTH = 800  # Max width in pixels for embedded images
MAX_IMAGE_HEIGHT = 600  # Max height in pixels
JPEG_QUALITY = 75  # JPEG quality for compression
BATCH_SIZE = 10  # Process images in batches for memory management


def optimize_image_for_word(image_path, max_width=MAX_IMAGE_WIDTH, max_height=MAX_IMAGE_HEIGHT):
    """
    Optimize an image for embedding in Word document.
    Returns a BytesIO object with the optimized image, or None if failed.
    """
    if not os.path.exists(image_path):
        return None
    
    try:
        with Image.open(image_path) as img:
            # Convert to RGB if necessary (handles PNG with transparency, etc.)
            if img.mode in ('RGBA', 'P', 'LA'):
                # Create white background for transparent images
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Calculate new dimensions while maintaining aspect ratio
            width, height = img.size
            if width > max_width or height > max_height:
                ratio = min(max_width / width, max_height / height)
                new_width = int(width * ratio)
                new_height = int(height * ratio)
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # Save to BytesIO
            buffer = io.BytesIO()
            img.save(buffer, format='JPEG', quality=JPEG_QUALITY, optimize=True)
            buffer.seek(0)
            return buffer
    except Exception as e:
        logger.warning(f"Error optimizing image {image_path}: {e}")
        return None


def set_rtl_paragraph(paragraph):
    """Set RTL direction for a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def add_rtl_paragraph(doc, text, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    """Add a right-to-left paragraph with Arabic text."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    set_rtl_paragraph(paragraph)
    
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    if bold:
        run.bold = True
    return paragraph


def add_section_title(doc, title):
    """Add a section title with page break before."""
    # Add page break
    doc.add_page_break()
    # Add title
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl_paragraph(paragraph)
    
    run = paragraph.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    # Add spacing
    doc.add_paragraph()


def add_decorative_separator(doc):
    """Add a decorative horizontal separator between sections."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create a table for the decorative bar
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.width = Inches(6)
    
    # Set cell background color (light gray)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D3D3D3')
    cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Set cell height
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '120')  # Small height for separator
    trPr.append(trHeight)
    
    # Remove borders
    for cell in table.rows[0].cells:
        cell._element.get_or_add_tcPr().append(
            OxmlElement('w:tcBorders')
        )
    
    doc.add_paragraph()  # Add spacing after separator


def add_image_to_doc(doc, image_path, width_inches=3.0, height_inches=None):
    """Add an optimized image to the document with proper sizing."""
    if not os.path.exists(image_path):
        return False
    
    try:
        # Try to use optimized image first
        optimized_buffer = optimize_image_for_word(image_path)
        
        if optimized_buffer:
            # Add optimized image
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(optimized_buffer, width=Inches(width_inches))
            optimized_buffer.close()
            return True
        
        # Fallback: try to add original image directly
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        return True
        
    except Exception as e:
        logger.warning(f"Error adding image {image_path}: {e}")
        return False


def add_optimized_image_to_cell(cell, image_path, width_inches=2.8):
    """Add an optimized image to a table cell."""
    if not image_path or not os.path.exists(image_path):
        return False
    
    try:
        optimized_buffer = optimize_image_for_word(image_path)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
        if optimized_buffer:
            run.add_picture(optimized_buffer, width=Inches(width_inches))
            optimized_buffer.close()
        else:
            # Fallback to original
            run.add_picture(image_path, width=Inches(width_inches))
        return True
    except Exception as e:
        logger.warning(f"Error adding image to cell: {e}")
        return False


def add_optimized_image_to_paragraph(run, image_path, width_inches=2.5):
    """Add an optimized image to a run in a paragraph."""
    if not image_path or not os.path.exists(image_path):
        return False
    
    try:
        optimized_buffer = optimize_image_for_word(image_path)
        
        if optimized_buffer:
            run.add_picture(optimized_buffer, width=Inches(width_inches))
            optimized_buffer.close()
        else:
            # Fallback to original
            run.add_picture(image_path, width=Inches(width_inches))
        return True
    except Exception as e:
        logger.warning(f"Error adding image to paragraph: {e}")
        return False


def generate_memory_book_word():
    """
    Generate a Word document (.docx) for the College Memory Book.
    Order: Meetings (latest first) → Colleagues → Memories (by category)
    """
    from api.models import MeetingCategory, Colleague, MemoryCategory
    
    # Create document
    doc = Document()
    
    # Set default paragraph style to RTL for the document
    # Note: python-docx doesn't directly support document-level RTL
    # We'll set it per paragraph
    
    # ============================================
    # 1️⃣ اللقاءات (Meetings) - Latest First
    # ============================================
    add_section_title(doc, "اللقاءات")
    
    # Get all meeting categories ordered by year descending (latest first)
    meeting_categories = MeetingCategory.objects.prefetch_related('photos').all().order_by('-year', '-created_at')
    
    for meeting_category in meeting_categories:
        photos = meeting_category.photos.all()
        if not photos.exists():
            continue
        
        # Meeting title - In RTL mode, LEFT alignment positions text on RIGHT side
        paragraph = doc.add_paragraph()
        
        # Set RTL direction first
        set_rtl_paragraph(paragraph)
        
        # In RTL mode: LEFT alignment = text appears on RIGHT side of page
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Explicitly set justification to LEFT (which appears on right in RTL)
        pPr = paragraph._element.get_or_add_pPr()
        # Remove any existing jc element
        for elem in pPr:
            if elem.tag.endswith('}jc'):
                pPr.remove(elem)
        # Add LEFT justification (appears on right in RTL mode)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')
        pPr.append(jc)
        
        run = paragraph.add_run(meeting_category.name)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Meeting description if exists
        if meeting_category.description:
            add_rtl_paragraph(doc, meeting_category.description, font_size=12)
        
        doc.add_paragraph()  # Spacing
        
        # Add photos in a grid (2 per row)
        photos_list = list(photos)
        for i in range(0, len(photos_list), 2):
            # Create table for 2 images side by side
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            for j, photo in enumerate(photos_list[i:i+2]):
                cell = table.rows[0].cells[j]
                cell.width = Inches(3.0)
                
                # Add optimized image to cell
                if photo.image:
                    add_optimized_image_to_cell(cell, photo.image.path, width_inches=2.8)
                
                # Add description if exists
                if photo.description_ar:
                    desc_para = cell.add_paragraph()
                    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    desc_run = desc_para.add_run(photo.description_ar)
                    desc_run.font.size = Pt(10)
                    desc_run.font.name = 'Arial'
                    desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            
            doc.add_paragraph()  # Spacing after row
        
        # Add spacing between meetings
        doc.add_paragraph()
        add_decorative_separator(doc)
    
    # Garbage collection after meetings section
    gc.collect()
    
    # ============================================
    # 2️⃣ الزملاء (Colleagues)
    # ============================================
    add_section_title(doc, "الزملاء")
    
    # Get all colleagues with photos
    colleagues = Colleague.objects.prefetch_related('archive_photos').all().order_by('name')
    
    for colleague in colleagues:
        # Check if colleague has any photos
        has_photo_1973 = colleague.photo_1973 and os.path.exists(colleague.photo_1973.path)
        has_latest_photo = colleague.latest_photo and os.path.exists(colleague.latest_photo.path)
        archive_photos = list(colleague.archive_photos.all())
        
        # Skip if no photos at all
        if not has_photo_1973 and not has_latest_photo and not archive_photos:
            continue
        
        # Colleague name
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl_paragraph(paragraph)
        run = paragraph.add_run(colleague.name)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Two images side-by-side at top: صورة 73 and الصورة الأخيرة
        if has_photo_1973 or has_latest_photo:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # صورة 73
            cell1 = table.rows[0].cells[0]
            cell1.width = Inches(3.0)
            para1 = cell1.paragraphs[0]
            para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl_paragraph(para1)
            label1 = para1.add_run("صورة 73")
            label1.font.size = Pt(12)
            label1.font.bold = True
            label1.font.name = 'Arial'
            label1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            
            if has_photo_1973:
                para1_img = cell1.add_paragraph()
                para1_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run1_img = para1_img.add_run()
                add_optimized_image_to_paragraph(run1_img, colleague.photo_1973.path, width_inches=2.5)
            
            # الصورة الأخيرة
            cell2 = table.rows[0].cells[1]
            cell2.width = Inches(3.0)
            para2 = cell2.paragraphs[0]
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl_paragraph(para2)
            label2 = para2.add_run("الصورة الأخيرة")
            label2.font.size = Pt(12)
            label2.font.bold = True
            label2.font.name = 'Arial'
            label2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            
            if has_latest_photo:
                para2_img = cell2.add_paragraph()
                para2_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run2_img = para2_img.add_run()
                add_optimized_image_to_paragraph(run2_img, colleague.latest_photo.path, width_inches=2.5)
            
            doc.add_paragraph()  # Spacing
        
        # Additional archive photos (max 2 per row, larger size)
        if archive_photos:
            for i in range(0, len(archive_photos), 2):
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                for j, archive_photo in enumerate(archive_photos[i:i+2]):
                    cell = table.rows[0].cells[j]
                    cell.width = Inches(3.5)
                    
                    if archive_photo.image:
                        add_optimized_image_to_cell(cell, archive_photo.image.path, width_inches=3.2)
                
                doc.add_paragraph()  # Spacing after row
        
        # Add decorative separator between colleagues
        doc.add_paragraph()
        add_decorative_separator(doc)
    
    # Garbage collection after colleagues section
    gc.collect()
    
    # ============================================
    # 3️⃣ الذكريات (Memories) - By Category
    # ============================================
    add_section_title(doc, "الذكريات")
    
    # Get all memory categories ordered by year descending (latest first)
    memory_categories = MemoryCategory.objects.prefetch_related('photos').all().order_by('-year', 'name')
    
    for memory_category in memory_categories:
        photos = memory_category.photos.all()
        if not photos.exists():
            continue
        
        # Category title with page break before
        doc.add_page_break()
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl_paragraph(paragraph)
        run = paragraph.add_run(memory_category.name)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Category description if exists
        if memory_category.description:
            add_rtl_paragraph(doc, memory_category.description, font_size=12)
        
        doc.add_paragraph()  # Spacing
        
        # Add photos in a grid (2 per row)
        photos_list = list(photos)
        for i in range(0, len(photos_list), 2):
            # Create table for 2 images side by side
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            for j, photo in enumerate(photos_list[i:i+2]):
                cell = table.rows[0].cells[j]
                cell.width = Inches(3.0)
                
                # Add optimized image to cell
                if photo.image:
                    add_optimized_image_to_cell(cell, photo.image.path, width_inches=2.8)
                
                # Add description if exists
                if photo.description_ar:
                    desc_para = cell.add_paragraph()
                    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    desc_run = desc_para.add_run(photo.description_ar)
                    desc_run.font.size = Pt(10)
                    desc_run.font.name = 'Arial'
                    desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            
            doc.add_paragraph()  # Spacing after row
        
        # Add spacing between categories
        doc.add_paragraph()
        add_decorative_separator(doc)
    
    # Force garbage collection to free memory
    gc.collect()
    
    return doc

